import os
import pandas as pd
from docx import Document
from datetime import datetime
from docx.shared import Pt, RGBColor  # Para definir tamanho, cor, etc.
import tkinter as tk
from tkinter import filedialog

BRANDING = "Pr&nche - by Fellype Rosa / Rastek Soluções"

def main():
    print(BRANDING)


def select_file(default_filename, title, filetypes):
    """
    Verifica se o arquivo 'default_filename' existe.
    Se não existir, abre uma caixa de diálogo para o usuário selecionar o arquivo.
    Retorna o caminho do arquivo selecionado.
    """
    if os.path.exists(default_filename):
        return default_filename
    else:
        root = tk.Tk()
        root.withdraw()  # Oculta a janela principal do tkinter
        file_path = filedialog.askopenfilename(title=title, filetypes=filetypes)
        root.destroy()  # Encerra a janela tkinter
        return file_path

def apply_formatting(run):
    """
    Aplica a formatação padrão:
      - Fonte: Arial
      - Tamanho: 12pt
      - Cor: Preta
      - Sem negrito
    """
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = False

def iter_paragraphs(doc):
    """
    Gera todos os parágrafos do documento, incluindo os que estão dentro de tabelas.
    """
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def substituir_runs(paragrafos, substituicoes):
    """
    Faz a substituição placeholder-por-valor em cada run dos parágrafos.
    Para cada run que contenha uma chave, substitui o texto e aplica a formatação padrão.
    """
    for paragrafo in paragrafos:
        for run in paragrafo.runs:
            for chave, valor in substituicoes.items():
                if chave in run.text:
                    run.text = run.text.replace(chave, valor)
                    apply_formatting(run)

def substituir_paragrafo_completo(paragrafo, substituicoes):
    """
    Se o parágrafo ainda contém algum placeholder, junta todo o texto, 
    realiza a substituição e recria o parágrafo com um único run com formatação padrão.
    """
    texto_completo = paragrafo.text
    texto_substituido = texto_completo
    for chave, valor in substituicoes.items():
        texto_substituido = texto_substituido.replace(chave, valor)
    
    if texto_substituido != texto_completo:
        p_element = paragrafo._element
        for child in list(p_element):
            p_element.remove(child)
        novo_run = paragrafo.add_run(texto_substituido)
        apply_formatting(novo_run)

def verificar_e_substituir_paragrafos(paragrafos, substituicoes):
    """
    Para cada parágrafo, se após a primeira passagem ainda houver algum placeholder,
    realiza a substituição considerando o parágrafo inteiro.
    """
    for paragrafo in paragrafos:
        texto = paragrafo.text
        for chave in substituicoes.keys():
            if chave in texto:
                substituir_paragrafo_completo(paragrafo, substituicoes)
                break

def verificar_chaves_nao_substituidas(paragrafos, substituicoes):
    """
    Retorna um conjunto com as chaves que ainda não foram substituídas em algum parágrafo.
    """
    chaves_nao_substituidas = set()
    for paragrafo in paragrafos:
        for chave in substituicoes.keys():
            if chave in paragrafo.text:
                chaves_nao_substituidas.add(chave)
    return chaves_nao_substituidas

def main():
    # ----- Parte 1: Seleção dos Arquivos Externos -----
    # Tenta encontrar os arquivos padrão; se não encontrar, solicita ao usuário.
    excel_file = select_file("Dados de Preenchimento.xlsx",  # Nome atualizado para o Excel
                             "Selecione a planilha de preenchimento",
                             [("Arquivos Excel", "*.xlsx")])
    if not excel_file:
        raise Exception("Planilha não selecionada. Encerrando.")

    modelo_file = select_file("MODELO - Oposição Administrativa - SKO Oyarzabal.docx",  # Nome atualizado para o modelo
                              "Selecione o documento modelo",
                              [("Documentos Word", "*.docx")])
    if not modelo_file:
        raise Exception("Documento modelo não selecionado. Encerrando.")

    # ----- Parte 2: Leitura do Excel e criação do dicionário de substituições -----
    df = pd.read_excel(excel_file)
    
    substituicoes = {
        row["Campo"]:
            (row["Valor a ser preenchido"].strftime("%d/%m/%Y")
             if isinstance(row["Valor a ser preenchido"], (pd.Timestamp, datetime))
             else str(row["Valor a ser preenchido"]))
        for index, row in df.iterrows()
        if pd.notna(row["Valor a ser preenchido"]) and str(row["Valor a ser preenchido"]).strip() != ""
    }
    
    # ----- Parte 3: Abrir o Documento Modelo e Obter Todos os Parágrafos -----
    doc = Document(modelo_file)
    all_paragraphs = list(iter_paragraphs(doc))
    
    # Primeira passagem: substituição run a run
    substituir_runs(all_paragraphs, substituicoes)
    
    # Segunda passagem: tratamento dos parágrafos com placeholders pendentes
    verificar_e_substituir_paragrafos(all_paragraphs, substituicoes)
    
    # (Opcional) Verificação final de placeholders não substituídos:
    # chaves_pendentes = verificar_chaves_nao_substituidas(all_paragraphs, substituicoes)
    
    # ----- Parte 4: Seleção do Local e Nome do Arquivo para Salvamento -----
    default_name_parts = []
    if '[nº processo terceiro]' in substituicoes:
        default_name_parts.append(substituicoes['[nº processo terceiro]'])
    if '[NOME CLIENTE]' in substituicoes:
        default_name_parts.append(substituicoes['[NOME CLIENTE]'])
    default_filename = " ".join(default_name_parts).strip() + ".docx" if default_name_parts else "Novo_Modelo.docx"
    
    root = tk.Tk()
    root.withdraw()
    
    save_path = filedialog.asksaveasfilename(
        title="Salvar arquivo como...",
        defaultextension=".docx",
        filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")],
        initialfile=default_filename
    )
    
    if save_path:
        doc.save(save_path)

if __name__ == "__main__":
    main()
